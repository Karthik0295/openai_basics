import os
from openai import OpenAI
from dotenv import load_dotenv
from fastapi.responses import JSONResponse
from fastapi import FastAPI, File, UploadFile, status
import fitz
import docx
import openpyxl
import tempfile
import io


load_dotenv()

client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),
)

def create_response(success: bool, payload=None, error=None, code=None):
    response = {
        "success": success,
        "payload": payload or {},
        "error": error or {}
    }
    return JSONResponse(content=response, status_code=code or (status.HTTP_200_OK if success else status.HTTP_400_BAD_REQUEST))

async def get_openai_response(message: str):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",  
        messages=[
            {"role": "system", "content": "You are a helpful assistant who has general knowledge information"},
            {"role": "user", "content": message}
        ],
        max_tokens=256,
        temperature=1,
    )
    return  response.choices[0].message.content.strip()

async def translator(language1: str, language2: str, message: str):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "system",
                "content": f"You are a helpful assistant who translates {language1} to {language2}."
            },
            {"role": "user", "content": message}
        ],
        max_tokens=256,
        temperature=1,
    )
    return response.choices[0].message.content.strip()


async def create_upload_file(file: UploadFile):
    file_extension = file.filename.split(".")[-1].lower()
    extracted_text = ""

    file_content = await file.read()

    if file_extension == "pdf":
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            extracted_text += page.get_text()
    
    elif file_extension == "docx":
        docx_file = await file.read()
        doc = docx.Document(io.BytesIO(docx_file))

    elif file_extension == "docx":
        
        with tempfile.NamedTemporaryFile(delete=True) as temp_file:
            temp_file.write(file_content)
            temp_file.flush()  
            doc = docx.Document(temp_file.name)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

    elif file_extension == "xlsx":
        
        with tempfile.NamedTemporaryFile(delete=True) as temp_file:
            temp_file.write(file_content)
            temp_file.flush()  
            wb = openpyxl.load_workbook(temp_file.name)
            sheet = wb.active
            for row in sheet.iter_rows(values_only=True):
                extracted_text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"

    else:
        return {"error": "Unsupported file type"}

    max_input_length = 4096  
    if len(extracted_text) > max_input_length:
        extracted_text = extracted_text[:max_input_length]

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant who summarizes text from the document."},
            {"role": "user", "content": extracted_text}
        ],
        max_tokens=256,
        temperature=0.7,
    )

    chat_response_doc = response.choices[0].message.content.strip()
    
    return { "summary": chat_response_doc}

async def transcribe_audio(file_path: str):
    
    with open(file_path, 'rb') as audio_file:
        response = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    
    return response.text

async def summarize_text(text: str):
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": f"Summarize the following text:\n\n{text}"}
        ],
        max_tokens=150,
        temperature=0.7,
    )
    
    return response.choices[0].message.content.strip()